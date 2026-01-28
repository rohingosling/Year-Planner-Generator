"""
Graph Paper section generator for Year Planner.

Creates graph paper pages with configurable grid dimensions and colors.
Each page has a grid on the recto with blank verso.
"""

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

from src.config import Config
from src.document import add_page_break, get_content_width, get_content_height, add_config_info_overlay
from src.utils.grid_image import generate_grid_image


# Target DPI for print quality
PRINT_DPI = 300

# Pixels per centimeter at target DPI
# 1 inch = 2.54 cm, so DPI / 2.54 = pixels per cm
PX_PER_CM = PRINT_DPI / 2.54

# Image output directory (relative to project root)
IMAGE_DIR = Path(__file__).parent.parent.parent / "assets" / "images"


def generate_graph_paper(document: Document, config: Config) -> None:
    """
    Generate the graph paper section.

    Creates configured number of graph paper pages. Each page has:
    - Grid on the recto (front)
    - Blank verso (back)

    Images are cached in assets/images/ and only regenerated if missing
    or configuration changes.

    Args:
        document: The Word document to add graph paper to.
        config: Configuration with graph paper settings.
    """
    # Get configuration values
    graph_config = config.raw.get('graph_paper', {})
    page_count = graph_config.get('page_count', 8)
    columns = graph_config.get('columns', 37)
    rows = graph_config.get('rows', 56)
    grid_color_percent = graph_config.get('grid_color_percent', 15)
    border_color_percent = graph_config.get('border_color_percent', 100)

    # Calculate content area dimensions (with gutter for binding)
    content_width_cm = get_content_width(config, include_gutter=True)
    content_height_cm = get_content_height(config)

    # Calculate image dimensions in pixels
    width_px = int(content_width_cm * PX_PER_CM)
    height_px = int(content_height_cm * PX_PER_CM)

    # Ensure image directory exists
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)

    # Generate filename based on configuration (enables caching)
    # Include pixel dimensions to ensure regeneration when page layout changes
    # (e.g., different margins result in different content area size)
    image_filename = (
        f"graph_paper_{columns}x{rows}_{grid_color_percent}_{border_color_percent}"
        f"_{width_px}x{height_px}px.png"
    )
    image_path = IMAGE_DIR / image_filename

    # Generate image only if it doesn't exist (cache)
    if not image_path.exists():
        generate_grid_image(
            width_px=width_px,
            height_px=height_px,
            columns=columns,
            rows=rows,
            grid_color_percent=grid_color_percent,
            border_color_percent=border_color_percent,
            output_path=str(image_path)
        )

    # Generate graph paper pages
    for page_num in range(page_count):
        # Insert image into document, sized to fill content area
        document.add_picture(
            str(image_path),
            width=Cm(content_width_cm),
            height=Cm(content_height_cm)
        )

        # Ensure the picture paragraph has no spacing
        # (document.add_picture creates a new paragraph for the picture)
        picture_para = document.paragraphs[-1]
        picture_para.paragraph_format.space_before = Pt(0)
        picture_para.paragraph_format.space_after = Pt(0)

        # Add overlay to grid page (recto)
        # Anchor to the picture's paragraph since the page is full
        add_config_info_overlay(document, config, is_recto=True,
                                anchor_paragraph=picture_para)

        # Add page break to create blank verso and position for next recto
        # Use minimize_height=True to prevent the page break paragraph from
        # affecting page layout
        page_break_para = add_page_break(document, minimize_height=True)

        # Add overlay to blank verso - anchor to page break paragraph
        # to avoid creating an extra paragraph that causes blank pages
        add_config_info_overlay(document, config, is_recto=False,
                                anchor_paragraph=page_break_para)
