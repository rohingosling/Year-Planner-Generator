"""
Style definitions for Year Planner document.

Provides consistent font and paragraph styling throughout the document.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.table import _Cell


# Font settings
FONT_NAME = "Times New Roman"
FONT_SIZE_TITLE = Pt(36)
FONT_SIZE_SUBTITLE = Pt(14)
FONT_SIZE_NORMAL = Pt(11)
FONT_SIZE_SMALL = Pt(9)

# Colors
COLOR_BLACK = RGBColor(0, 0, 0)


def apply_title_style(paragraph: Paragraph) -> None:
    """
    Apply title styling to a paragraph.

    Args:
        paragraph: The paragraph to style.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_TITLE
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK


def apply_subtitle_style(paragraph: Paragraph) -> None:
    """
    Apply subtitle styling to a paragraph.

    Args:
        paragraph: The paragraph to style.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_SUBTITLE
    run.font.bold = False
    run.font.color.rgb = COLOR_BLACK


def apply_normal_style(paragraph: Paragraph) -> None:
    """
    Apply normal text styling to a paragraph.

    Args:
        paragraph: The paragraph to style.
    """
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.color.rgb = COLOR_BLACK


def apply_cell_style(cell: _Cell, bold: bool = False,
                     alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """
    Apply styling to a table cell.

    Args:
        cell: The table cell to style.
        bold: Whether to make the text bold.
        alignment: Text alignment within the cell.
    """
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.bold = bold
    run.font.color.rgb = COLOR_BLACK
